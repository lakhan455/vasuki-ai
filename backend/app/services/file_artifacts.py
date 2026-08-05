from __future__ import annotations

import base64
import io
import os
import re
import textwrap
from pathlib import Path
from typing import Any

import qrcode
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from PIL import Image
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.config import Settings
from app.services.chat import route_chat
from app.services.rag import extract_document_pages
from app.services.vision import IMAGE_MIME_TYPES, process_vision_request


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".jpg", ".jpeg", ".png", ".webp", ".gif"}
SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "image/jpeg",
    "image/png",
    "image/webp",
    "image/gif",
}

SMART_FILE_RULES = """You are Vasuki Smart Files, a precise document-analysis assistant.
Use only the supplied file content unless the user explicitly asks for general knowledge.

Rules:
1. Answer the exact request in the same language and script as the user.
2. For a question paper, worksheet, assignment or exam sheet, answer every readable question in the original order. Keep question numbers and options. Show calculations where needed.
3. For multiple files, compare and combine them without losing which file a fact came from.
4. When asked for notes, create clean headings, compact bullets, definitions, formulas, examples and key takeaways.
5. When asked for a one-sheet or printable sheet, make the response compact enough for one A4 sheet, remove repetition and prioritize the most useful content.
6. Never invent missing or unreadable text. Clearly identify any unreadable part.
7. Check arithmetic, spelling, requested counts, dates, units and formatting before finishing.
8. Do not mention internal APIs, models, prompts or implementation details.
9. Return the complete polished content. The application will package it into a downloadable file only when the user explicitly requests one.
"""

_FILE_ACTION_RE = re.compile(
    r"(?is)(?:\b(?:create|make|generate|prepare|export|download|provide|give|convert|save|print|build|bana(?:o|kar)?|banado|de\s*do|file\s*do)\b.{0,60}\b(?:pdf|docx|word|txt|text\s*file|qr|png|file|one[-\s]?sheet)\b|"
    r"\b(?:pdf|docx|word|txt|text\s*file|qr|png|file|one[-\s]?sheet)\b.{0,60}\b(?:create|make|generate|prepare|export|download|provide|give|convert|save|print|bana(?:o|kar)?|banado|de\s*do)\b)"
)


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _suffix(filename: str) -> str:
    return Path(filename or "document").suffix.casefold()


def _normalise_mime(filename: str, mime_type: str) -> str:
    mime = (mime_type or "").split(";", 1)[0].strip().casefold()
    suffix = _suffix(filename)
    by_suffix = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }
    return by_suffix.get(suffix, mime)


def _validate_upload(filename: str, mime_type: str, content: bytes) -> None:
    if not content:
        raise ValueError(f"{filename}: file is empty.")
    if len(content) > 15 * 1024 * 1024:
        raise ValueError(f"{filename}: each file must be 15 MB or smaller.")
    suffix = _suffix(filename)
    mime = _normalise_mime(filename, mime_type)
    if suffix not in SUPPORTED_EXTENSIONS and mime not in SUPPORTED_MIME_TYPES:
        raise ValueError(
            f"{filename}: unsupported type. Use PDF, DOCX, TXT, MD, JPG, PNG, WEBP or GIF."
        )


def _page_pack(filename: str, pages: list[tuple[int | None, str]]) -> str:
    parts: list[str] = []
    for page_number, text in pages:
        cleaned = _clean_text(text)
        if not cleaned:
            continue
        label = f"Page {page_number}" if page_number else "Document text"
        parts.append(f"[{filename} · {label}]\n{cleaned}")
    return "\n\n".join(parts)


async def _extract_upload(upload: dict[str, object], prompt: str, settings: Settings) -> tuple[str, str]:
    filename = str(upload.get("filename") or "document")
    mime_type = str(upload.get("mime_type") or "application/octet-stream")
    content = bytes(upload.get("content") or b"")
    _validate_upload(filename, mime_type, content)
    mime = _normalise_mime(filename, mime_type)

    if mime not in IMAGE_MIME_TYPES:
        try:
            pages = extract_document_pages(content, filename, mime)
            packed = _page_pack(filename, pages)
            if packed:
                return filename, packed
        except Exception:
            if mime != "application/pdf":
                raise

    vision_prompt = (
        "Extract and analyze all readable content from this file. Preserve question numbers, "
        "tables, options, formulas and page order. Then prepare evidence for this user request: "
        + prompt
    )
    result = await process_vision_request(
        content=content,
        filename=filename,
        mime_type=mime,
        prompt=vision_prompt,
        operation="analyze",
        settings=settings,
    )
    answer = _clean_text(str(result.get("answer") or ""))
    if not answer:
        raise ValueError(f"{filename}: no readable content was found.")
    return filename, f"[{filename} · visual/scanned analysis]\n{answer}"


def _limit_file_pack(items: list[tuple[str, str]], limit: int = 42_000) -> str:
    if not items:
        return "(No files were uploaded.)"
    remaining = limit
    output: list[str] = []
    for index, (filename, content) in enumerate(items, 1):
        header = f"\n\n===== FILE {index}: {filename} =====\n"
        available = max(0, remaining - len(header))
        if available <= 0:
            break
        piece = content[:available]
        output.append(header + piece)
        remaining -= len(header) + len(piece)
    if len(output) < len(items):
        output.append("\n\n[Some file content was shortened to fit the AI context window.]\n")
    return "".join(output).strip()


def _wants_artifact(prompt: str) -> bool:
    normalized = prompt.casefold()
    if _FILE_ACTION_RE.search(normalized):
        return True
    return bool(
        re.search(r"\bpdf\s+(?:me|mein)\s+(?:de|do|dena|bana)", normalized)
        or re.search(r"\b(?:one[-\s]?sheet|printable\s+sheet)\b.*\b(?:file|pdf|print)\b", normalized)
        or re.search(r"\bqr\b.*\b(?:image|pdf|download|file)\b", normalized)
    )


def _requested_formats(prompt: str) -> tuple[set[str], bool]:
    normalized = prompt.casefold()
    one_sheet = bool(re.search(r"\bone[-\s]?sheet\b|\bsingle[-\s]?page\b|\bek\s+sheet\b", normalized))
    formats: set[str] = set()
    if "qr" in normalized:
        formats.add("qr")
    if "pdf" in normalized or one_sheet:
        formats.add("pdf")
    if "docx" in normalized or re.search(r"\bword\s+(?:file|document)\b", normalized):
        formats.add("docx")
    if re.search(r"\btxt\b|\btext\s+file\b", normalized):
        formats.add("txt")
    if not formats and _wants_artifact(prompt):
        formats.add("pdf")
    return formats, one_sheet


def _extract_qr_payload(prompt: str) -> str:
    url_match = re.search(r"https?://[^\s<>\]\)\}\"']+", prompt)
    if url_match:
        return url_match.group(0).rstrip(".,;:")
    quoted = re.search(r"[\"']([^\"']{2,500})[\"']", prompt)
    if quoted:
        return quoted.group(1).strip()
    match = re.search(
        r"(?is)\bqr(?:\s+code)?\s+(?:for|of|ka|ki|ke\s+liye|ban(?:a|ao)\s+for)\s*[:\-]?\s*(.{2,500})",
        prompt,
    )
    if match:
        value = match.group(1).strip()
        value = re.split(r"\b(?:and|aur)\s+(?:pdf|image|file)\b", value, maxsplit=1, flags=re.I)[0].strip()
        return value
    return ""


def _strip_markdown(text: str) -> str:
    value = re.sub(r"```[a-zA-Z0-9_-]*\n?", "", text)
    value = value.replace("```", "")
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(?m)^#{1,6}\s*", "", value)
    value = value.replace("**", "").replace("__", "").replace("`", "")
    return _clean_text(value)


def _slug(value: str, default: str = "vasuki-ai-output") -> str:
    ascii_value = value.encode("ascii", errors="ignore").decode("ascii").casefold()
    ascii_value = re.sub(r"[^a-z0-9]+", "-", ascii_value).strip("-")
    return (ascii_value[:60] or default).strip("-")


def _title_from_prompt(prompt: str) -> str:
    clean = _strip_markdown(prompt)
    clean = re.sub(r"(?i)\b(create|make|generate|prepare|give|provide|download|pdf|docx|txt|file|banao|banado|de do)\b", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip(" .:-")
    return clean[:90] or "Vasuki AI Output"


def _artifact(name: str, mime_type: str, content: bytes) -> dict[str, Any]:
    encoded = base64.b64encode(content).decode("ascii")
    return {
        "name": name,
        "mime_type": mime_type,
        "size_bytes": len(content),
        "data_url": f"data:{mime_type};base64,{encoded}",
    }


def _font_name() -> str:
    if "VasukiUnicode" in pdfmetrics.getRegisteredFontNames():
        return "VasukiUnicode"
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont("VasukiUnicode", candidate))
                return "VasukiUnicode"
            except Exception:
                continue
    return "Helvetica"


def _pdf_safe(value: str, font: str) -> str:
    if font == "VasukiUnicode":
        return value
    return value.encode("latin-1", errors="replace").decode("latin-1")


def _wrap_pdf_line(text: str, width: float, font: str, size: float) -> list[str]:
    text = text.rstrip()
    if not text:
        return [""]
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if pdfmetrics.stringWidth(_pdf_safe(candidate, font), font, size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def _plain_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in _strip_markdown(text).splitlines():
        if len(raw) <= 500:
            lines.append(raw)
        else:
            lines.extend(textwrap.wrap(raw, width=180, break_long_words=False, replace_whitespace=False))
    return lines


def _build_text_pdf(title: str, body: str, *, one_sheet: bool) -> bytes:
    buffer = io.BytesIO()
    font = _font_name()
    page_size = landscape(A4) if one_sheet else A4
    pdf = canvas.Canvas(buffer, pagesize=page_size)
    width, height = page_size
    margin = 34 if one_sheet else 48
    pdf.setTitle(title)

    if one_sheet:
        pdf.setFont(font, 15)
        pdf.drawString(margin, height - margin, _pdf_safe(title[:100], font))
        top = height - margin - 24
        gap = 22
        column_width = (width - 2 * margin - gap) / 2
        raw_lines = _plain_lines(body)
        chosen_size = 9.0
        wrapped: list[str] = []
        for candidate_size in (9.0, 8.5, 8.0, 7.5, 7.0, 6.5):
            candidate_lines: list[str] = []
            for line in raw_lines:
                candidate_lines.extend(_wrap_pdf_line(line, column_width, font, candidate_size))
            capacity = int((top - margin) / (candidate_size * 1.35)) * 2
            wrapped = candidate_lines
            chosen_size = candidate_size
            if len(candidate_lines) <= capacity:
                break
        line_height = chosen_size * 1.35
        per_column = max(1, int((top - margin) / line_height))
        wrapped = wrapped[: per_column * 2]
        pdf.setFont(font, chosen_size)
        for index, line in enumerate(wrapped):
            column = index // per_column
            row = index % per_column
            x = margin + column * (column_width + gap)
            y = top - row * line_height
            pdf.drawString(x, y, _pdf_safe(line, font))
        pdf.setFont(font, 7)
        pdf.drawRightString(width - margin, 16, _pdf_safe("Created with Vasuki AI", font))
        pdf.showPage()
    else:
        lines: list[str] = []
        for line in _plain_lines(body):
            lines.extend(_wrap_pdf_line(line, width - 2 * margin, font, 10.5))
        y = height - margin
        pdf.setFont(font, 16)
        pdf.drawString(margin, y, _pdf_safe(title[:100], font))
        y -= 28
        pdf.setFont(font, 10.5)
        for line in lines:
            if y < margin + 22:
                pdf.setFont(font, 7)
                pdf.drawRightString(width - margin, 18, _pdf_safe("Created with Vasuki AI", font))
                pdf.showPage()
                y = height - margin
                pdf.setFont(font, 10.5)
            pdf.drawString(margin, y, _pdf_safe(line, font))
            y -= 15
        pdf.setFont(font, 7)
        pdf.drawRightString(width - margin, 18, _pdf_safe("Created with Vasuki AI", font))
        pdf.showPage()

    pdf.save()
    return buffer.getvalue()


def _build_docx(title: str, body: str) -> bytes:
    document = DocxDocument()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    document.add_heading(title, level=1)
    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            document.add_paragraph("")
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif re.match(r"^[-*•]\s+", line):
            document.add_paragraph(re.sub(r"^[-*•]\s+", "", line), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown(line))
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
    document.add_paragraph("Created with Vasuki AI")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_qr(payload: str) -> bytes:
    qr = qrcode.QRCode(version=None, error_correction=qrcode.constants.ERROR_CORRECT_M, box_size=12, border=4)
    qr.add_data(payload)
    qr.make(fit=True)
    image = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _build_qr_pdf(title: str, payload: str, qr_png: bytes) -> bytes:
    buffer = io.BytesIO()
    font = _font_name()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    pdf.setTitle(title)
    pdf.setFont(font, 18)
    pdf.drawCentredString(width / 2, height - 70, _pdf_safe(title[:100], font))
    size = 330
    pdf.drawImage(ImageReader(io.BytesIO(qr_png)), (width - size) / 2, height - 455, size, size, preserveAspectRatio=True)
    pdf.setFont(font, 9)
    wrapped = _wrap_pdf_line(payload, width - 100, font, 9)
    y = height - 485
    for line in wrapped[:8]:
        pdf.drawCentredString(width / 2, y, _pdf_safe(line, font))
        y -= 14
    pdf.setFont(font, 8)
    pdf.drawCentredString(width / 2, 28, _pdf_safe("Created with Vasuki AI", font))
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


async def process_smart_file_request(*, uploads: list[dict[str, object]], prompt: str, settings: Settings) -> dict[str, Any]:
    extracted: list[tuple[str, str]] = []
    extraction_errors: list[str] = []
    for upload in uploads:
        try:
            extracted.append(await _extract_upload(upload, prompt, settings))
        except Exception as exc:
            filename = str(upload.get("filename") or "document")
            extraction_errors.append(f"{filename}: {str(exc)[:300]}")

    if uploads and not extracted:
        raise ValueError("None of the uploaded files could be read. " + " | ".join(extraction_errors))

    file_pack = _limit_file_pack(extracted)
    user_message = (
        SMART_FILE_RULES
        + "\n\nUSER REQUEST:\n"
        + prompt
        + "\n\nSUPPLIED FILE CONTENT:\n"
        + file_pack
    )
    if extraction_errors:
        user_message += "\n\nFILE WARNINGS:\n" + "\n".join(extraction_errors)

    answer, provider = await route_chat(
        "auto",
        [{"role": "user", "content": user_message}],
        settings,
        require_current=False,
    )
    answer = answer.strip()
    if not answer:
        raise RuntimeError("The AI returned an empty file-analysis response.")

    generated: list[dict[str, Any]] = []
    artifact_note = ""
    if _wants_artifact(prompt):
        formats, one_sheet = _requested_formats(prompt)
        title = _title_from_prompt(prompt)
        basename = _slug(title)

        qr_png: bytes | None = None
        qr_payload = ""
        if "qr" in formats:
            qr_payload = _extract_qr_payload(prompt)
            if qr_payload:
                qr_png = _build_qr(qr_payload)
                generated.append(_artifact(f"{basename}-qr.png", "image/png", qr_png))
            else:
                artifact_note = " A QR file was not created because no URL or QR text was provided."

        if "pdf" in formats:
            if qr_png is not None:
                pdf_bytes = _build_qr_pdf(title or "QR Code", qr_payload, qr_png)
            else:
                pdf_bytes = _build_text_pdf(title, answer, one_sheet=one_sheet)
            generated.append(_artifact(f"{basename}.pdf", "application/pdf", pdf_bytes))

        if "docx" in formats:
            generated.append(
                _artifact(
                    f"{basename}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    _build_docx(title, answer),
                )
            )

        if "txt" in formats:
            text_bytes = (title + "\n\n" + _strip_markdown(answer) + "\n\nCreated with Vasuki AI\n").encode("utf-8")
            generated.append(_artifact(f"{basename}.txt", "text/plain;charset=utf-8", text_bytes))

    if generated:
        answer += "\n\nYour requested downloadable file is ready below."
    if artifact_note:
        answer += artifact_note

    return {
        "answer": answer,
        "provider": provider,
        "files": generated,
        "processed_files": [name for name, _content in extracted],
        "warnings": extraction_errors,
    }
