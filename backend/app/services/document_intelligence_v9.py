from __future__ import annotations

import asyncio
import difflib
import re
from io import BytesIO
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.config import Settings
from app.services.chat import route_chat
from app.services.vision import IMAGE_MIME_TYPES, process_vision_request


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".gif",
}

_TOKEN_RE = re.compile(r"[A-Za-z0-9\u0900-\u097F]{2,}")
_CITATION_RE = re.compile(r"\[([A-Z]\d+:(?:P\d+|S\d+|L\d+)(?::B\d+)?)\]")


def _ext(name: str) -> str:
    index = str(name or "").rfind(".")
    return str(name or "")[index:].casefold() if index >= 0 else ""


def _tokens(text: str) -> set[str]:
    return {item.casefold() for item in _TOKEN_RE.findall(str(text or ""))}


def _chunks(text: str, max_chars: int = 1800) -> list[str]:
    clean = str(text or "").replace("\x00", "").strip()
    if not clean:
        return []
    paragraphs = [re.sub(r"\s+", " ", item).strip() for item in re.split(r"\n\s*\n", clean)]
    paragraphs = [item for item in paragraphs if item]
    output: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            if current:
                output.append(current)
                current = ""
            for start in range(0, len(paragraph), max_chars):
                output.append(paragraph[start:start + max_chars])
            continue
        candidate = paragraph if not current else current + "\n\n" + paragraph
        if len(candidate) > max_chars:
            output.append(current)
            current = paragraph
        else:
            current = candidate
    if current:
        output.append(current)
    return output


def _block(
    *,
    source_id: str,
    location: str,
    block_index: int,
    text: str,
    page: int | None = None,
    section: str | None = None,
    kind: str = "text",
) -> dict[str, Any]:
    citation_id = f"{source_id}:{location}:B{block_index}"
    return {
        "citation_id": citation_id,
        "source_id": source_id,
        "page": page,
        "section": section,
        "kind": kind,
        "text": text,
        "word_count": len(text.split()),
    }


def extract_text_document(content: bytes, *, filename: str, source_id: str) -> dict[str, Any]:
    decoded = content.decode("utf-8", errors="replace")
    lines = decoded.splitlines()
    blocks: list[dict[str, Any]] = []
    current_lines: list[str] = []
    start_line = 1
    block_index = 1
    for line_number, line in enumerate(lines, 1):
        current_lines.append(line)
        joined = "\n".join(current_lines)
        if len(joined) >= 1600:
            blocks.append(
                _block(
                    source_id=source_id,
                    location=f"L{start_line}",
                    block_index=block_index,
                    text=joined.strip(),
                    section=f"Lines {start_line}-{line_number}",
                )
            )
            current_lines = []
            start_line = line_number + 1
            block_index += 1
    if current_lines:
        end_line = max(start_line, len(lines))
        blocks.append(
            _block(
                source_id=source_id,
                location=f"L{start_line}",
                block_index=block_index,
                text="\n".join(current_lines).strip(),
                section=f"Lines {start_line}-{end_line}",
            )
        )
    blocks = [item for item in blocks if item["text"]]
    return {
        "source_id": source_id,
        "name": filename,
        "type": "text",
        "pages": None,
        "blocks": blocks,
        "warnings": [],
    }


def extract_pdf_document(content: bytes, *, filename: str, source_id: str) -> dict[str, Any]:
    reader = PdfReader(BytesIO(content))
    blocks: list[dict[str, Any]] = []
    warnings: list[str] = []
    for page_number, page in enumerate(reader.pages, 1):
        text = str(page.extract_text() or "").strip()
        if not text:
            warnings.append(f"{filename}: page {page_number} has no embedded text.")
            continue
        for block_index, chunk in enumerate(_chunks(text), 1):
            blocks.append(
                _block(
                    source_id=source_id,
                    location=f"P{page_number}",
                    block_index=block_index,
                    text=chunk,
                    page=page_number,
                    section=f"Page {page_number}",
                )
            )
    return {
        "source_id": source_id,
        "name": filename,
        "type": "pdf",
        "pages": len(reader.pages),
        "blocks": blocks,
        "warnings": warnings,
    }


def extract_docx_document(content: bytes, *, filename: str, source_id: str) -> dict[str, Any]:
    document = Document(BytesIO(content))
    blocks: list[dict[str, Any]] = []
    section = "Document"
    section_number = 1
    block_index = 1
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal block_index, paragraph_buffer
        text = "\n".join(paragraph_buffer).strip()
        if text:
            for chunk in _chunks(text):
                blocks.append(
                    _block(
                        source_id=source_id,
                        location=f"S{section_number}",
                        block_index=block_index,
                        text=chunk,
                        section=section,
                    )
                )
                block_index += 1
        paragraph_buffer = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = str(getattr(paragraph.style, "name", "") or "").casefold()
        if style.startswith("heading"):
            flush()
            section_number += 1
            section = text
            continue
        paragraph_buffer.append(text)
    flush()

    for table_index, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(item for item in rows if item.strip())
        if text:
            blocks.append(
                _block(
                    source_id=source_id,
                    location=f"S{section_number + table_index}",
                    block_index=block_index,
                    text=text,
                    section=f"Table {table_index}",
                    kind="table",
                )
            )
            block_index += 1

    return {
        "source_id": source_id,
        "name": filename,
        "type": "docx",
        "pages": None,
        "blocks": blocks,
        "warnings": [],
    }


async def extract_image_document(
    content: bytes,
    *,
    filename: str,
    mime_type: str,
    source_id: str,
    settings: Settings,
) -> dict[str, Any]:
    result = await process_vision_request(
        content=content,
        filename=filename,
        mime_type=mime_type,
        prompt=(
            "OCR V2: Transcribe every readable word, number, heading, label and table "
            "from this image. Preserve reading order and line/section structure. "
            "Do not guess unreadable text. Return transcription only."
        ),
        operation="analyze",
        settings=settings,
    )
    text = str(result.get("answer") or "").strip()
    blocks = [
        _block(
            source_id=source_id,
            location="P1",
            block_index=index,
            text=chunk,
            page=1,
            section="OCR page 1",
            kind="ocr",
        )
        for index, chunk in enumerate(_chunks(text), 1)
    ]
    return {
        "source_id": source_id,
        "name": filename,
        "type": "image-ocr",
        "pages": 1,
        "blocks": blocks,
        "warnings": [],
        "ocr_provider": result.get("provider"),
    }


async def extract_upload(
    upload: dict[str, Any],
    *,
    source_id: str,
    settings: Settings,
) -> dict[str, Any]:
    filename = str(upload.get("filename") or "document")
    content = bytes(upload.get("content") or b"")
    mime_type = str(upload.get("mime_type") or "application/octet-stream").split(";", 1)[0].casefold()
    extension = _ext(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"{filename}: unsupported document type.")

    if extension == ".pdf" or mime_type == "application/pdf":
        document = extract_pdf_document(content, filename=filename, source_id=source_id)
        if not document["blocks"]:
            fallback = await process_vision_request(
                content=content,
                filename=filename,
                mime_type="application/pdf",
                prompt=(
                    "OCR V2: This PDF has little or no embedded text. Transcribe all "
                    "readable content faithfully. Preserve page markers when visible "
                    "and never invent unreadable text."
                ),
                operation="analyze",
                settings=settings,
            )
            text = str(fallback.get("answer") or "").strip()
            document["blocks"] = [
                _block(
                    source_id=source_id,
                    location="S1",
                    block_index=index,
                    text=chunk,
                    section="Vision OCR fallback",
                    kind="ocr",
                )
                for index, chunk in enumerate(_chunks(text), 1)
            ]
            document["ocr_provider"] = fallback.get("provider")
            document["warnings"].append(
                f"{filename}: page-level citations are unavailable for the vision OCR fallback."
            )
        return document

    if extension == ".docx":
        return extract_docx_document(content, filename=filename, source_id=source_id)

    if extension in {".txt", ".md"}:
        return extract_text_document(content, filename=filename, source_id=source_id)

    if mime_type in IMAGE_MIME_TYPES or extension in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
        if not mime_type.startswith("image/"):
            mime_type = {
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".png": "image/png",
                ".webp": "image/webp",
                ".gif": "image/gif",
            }.get(extension, "image/jpeg")
        return await extract_image_document(
            content,
            filename=filename,
            mime_type=mime_type,
            source_id=source_id,
            settings=settings,
        )

    raise ValueError(f"{filename}: unsupported document type.")


async def extract_uploads(
    uploads: list[dict[str, Any]],
    *,
    settings: Settings,
) -> dict[str, Any]:
    documents = []
    warnings = []
    for index, upload in enumerate(uploads, 1):
        document = await extract_upload(upload, source_id=f"D{index}", settings=settings)
        documents.append(document)
        warnings.extend(document.get("warnings") or [])
    return {
        "documents": documents,
        "warnings": warnings,
        "total_blocks": sum(len(item.get("blocks") or []) for item in documents),
    }


def flatten_blocks(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for document in documents:
        for block in document.get("blocks") or []:
            rows.append({
                **block,
                "document_name": document.get("name"),
                "document_type": document.get("type"),
            })
    return rows


def select_evidence(query: str, blocks: list[dict[str, Any]], limit: int = 12) -> list[dict[str, Any]]:
    query_tokens = _tokens(query)
    scored = []
    for order, block in enumerate(blocks):
        text_tokens = _tokens(str(block.get("text") or ""))
        overlap = len(query_tokens & text_tokens)
        coverage = overlap / max(1, len(query_tokens))
        density = overlap / max(1, min(80, len(text_tokens)))
        score = coverage * 4.0 + density + (0.0001 * (len(blocks) - order))
        scored.append((score, block))
    scored.sort(key=lambda item: item[0], reverse=True)
    if query_tokens and any(score > 0.001 for score, _ in scored):
        return [block for score, block in scored[: max(1, limit)] if score > 0.001][:limit]
    return [block for _score, block in scored[:limit]]


def _citation_record(block: dict[str, Any]) -> dict[str, Any]:
    return {
        "citation_id": block.get("citation_id"),
        "document": block.get("document_name"),
        "page": block.get("page"),
        "section": block.get("section"),
        "kind": block.get("kind"),
        "excerpt": str(block.get("text") or "")[:420],
    }


async def answer_with_citations(
    *,
    prompt: str,
    documents: list[dict[str, Any]],
    settings: Settings,
    evidence_limit: int = 14,
) -> dict[str, Any]:
    blocks = flatten_blocks(documents)
    evidence = select_evidence(prompt, blocks, limit=evidence_limit)
    if not evidence:
        return {
            "answer": "No readable document evidence was found.",
            "provider": None,
            "citations": [],
            "evidence": [],
        }
    source_text = "\n\n".join(
        f"[{item['citation_id']}] {item.get('document_name')} "
        f"(page={item.get('page')}, section={item.get('section')}):\n{item.get('text')}"
        for item in evidence
    )
    messages = [{
        "role": "user",
        "content": (
            "You are Vasuki Document Intelligence V3. Answer only from the supplied "
            "document evidence. Cite every important factual claim with the exact source "
            "ID in square brackets, for example [D1:P2:B1]. If the evidence does not "
            "support something, say so. Preserve the user's language.\n\n"
            f"USER REQUEST:\n{prompt.strip()}\n\nDOCUMENT EVIDENCE:\n{source_text[:42000]}"
        ),
    }]
    answer, provider = await route_chat("auto", messages, settings, require_current=False)
    cited = []
    seen = set()
    evidence_by_id = {str(item.get("citation_id")): item for item in evidence}
    for match in _CITATION_RE.finditer(str(answer or "")):
        citation_id = match.group(1)
        if citation_id in seen or citation_id not in evidence_by_id:
            continue
        seen.add(citation_id)
        cited.append(_citation_record(evidence_by_id[citation_id]))
    return {
        "answer": answer,
        "provider": provider,
        "citations": cited,
        "evidence": [_citation_record(item) for item in evidence],
    }


def deterministic_compare(documents: list[dict[str, Any]]) -> dict[str, Any]:
    if len(documents) < 2:
        raise ValueError("At least two documents are required for comparison.")
    left, right = documents[0], documents[1]
    left_blocks = [str(item.get("text") or "").strip() for item in left.get("blocks") or []]
    right_blocks = [str(item.get("text") or "").strip() for item in right.get("blocks") or []]
    left_text = "\n".join(left_blocks)[:100000]
    right_text = "\n".join(right_blocks)[:100000]
    similarity = difflib.SequenceMatcher(None, left_text, right_text, autojunk=False).ratio()

    left_norm = {re.sub(r"\s+", " ", text).casefold(): text for text in left_blocks if text}
    right_norm = {re.sub(r"\s+", " ", text).casefold(): text for text in right_blocks if text}
    removed = [left_norm[key][:700] for key in left_norm.keys() - right_norm.keys()][:12]
    added = [right_norm[key][:700] for key in right_norm.keys() - left_norm.keys()][:12]
    return {
        "left": left.get("name"),
        "right": right.get("name"),
        "similarity_percent": round(similarity * 100, 2),
        "added_samples": added,
        "removed_samples": removed,
    }


async def compare_with_citations(
    *,
    prompt: str,
    documents: list[dict[str, Any]],
    settings: Settings,
) -> dict[str, Any]:
    deterministic = deterministic_compare(documents)
    request = prompt.strip() or (
        "Compare these documents. Explain the most important similarities, differences, "
        "changed facts, changed numbers and missing/new sections."
    )
    answer = await answer_with_citations(
        prompt=request,
        documents=documents[:4],
        settings=settings,
        evidence_limit=20,
    )
    return {
        "comparison": deterministic,
        **answer,
    }
